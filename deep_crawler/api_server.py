#!/usr/bin/env python3
import os
import sys
import threading
import queue
import asyncio
import hashlib
import uuid
import toml
from datetime import datetime
from io import BytesIO
from pathlib import Path
from flask import Flask, request, jsonify, Response, send_file
from flask_cors import CORS
import markdown
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Add the parent directory to the path so we can import deep_crawler
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Try to import enhanced CLI, fallback to traditional
try:
    from deep_crawler.enhanced_cli import main as cli_main
    print("🚀 API Server: Using LangChain Enhanced CLI")
except ImportError as e:
    print(f"⚠️ Enhanced CLI not available: {e}")
    from deep_crawler.cli import main as cli_main
    print("🔄 API Server: Using Traditional CLI")

from deep_crawler import reports_db

# Load configuration
CONFIG = toml.load(Path(__file__).parent.parent / "config.toml")

app = Flask(__name__)
CORS(app)

# No longer need in-memory storage - using database instead

class StreamCapture:
    def __init__(self):
        self.queue = queue.Queue()
        self.finished = False
        self.content = []
    
    def write(self, text):
        if text.strip():
            self.queue.put(text)
            self.content.append(text)
    
    def flush(self):
        pass
    
    def finish(self):
        self.finished = True
    
    def get_full_content(self):
        return '\n'.join(self.content)

def convert_to_docx(markdown_content, title):
    """Convert markdown content to DOCX format"""
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(title, 0)
    
    # Convert markdown to HTML first, then extract text
    html = markdown.markdown(markdown_content)
    
    # Simple conversion - split by lines and handle basic formatting
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def convert_to_pdf(markdown_content, title):
    """Convert markdown content to PDF format"""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    # Process markdown content
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 12))
            continue
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], styles['Heading1']))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading3']))
        else:
            story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 6))
    
    doc.build(story)
    output.seek(0)
    return output

@app.route('/api/research', methods=['GET'])
def research():
    question = request.args.get('q')
    if not question or not question.strip():
        return jsonify({'error': 'Missing or empty q parameter'}), 400
    
    # Generate a unique ID for this research session
    research_id = str(uuid.uuid4())
    
    def generate():
        # Create a custom stdout capture
        capture = StreamCapture()
        original_stdout = sys.stdout
        
        try:
            # Replace stdout with our capture
            sys.stdout = capture
            
            # Run the CLI function in a separate thread
            def run_cli():
                try:
                    markdown_content = cli_main(question)
                    print(f"CLI returned content of length: {len(markdown_content) if markdown_content else 0}")
                    # Store the generated content in database
                    reports_db.store_report(
                        research_id=research_id,
                        question=question,
                        content=markdown_content,
                        stream_output=capture.get_full_content()
                    )
                    print(f"Stored report in database: {research_id}")
                except Exception as e:
                    print(f"CLI error: {str(e)}")
                    capture.queue.put(f"ERROR: {str(e)}")
                    reports_db.store_report(
                        research_id=research_id,
                        question=question,
                        error=str(e),
                        stream_output=capture.get_full_content()
                    )
                finally:
                    capture.finish()
            
            cli_thread = threading.Thread(target=run_cli)
            cli_thread.start()
            
            # Send the research ID first
            yield f"data: RESEARCH_ID:{research_id}\n\n"
            
            # Stream the output
            while cli_thread.is_alive() or not capture.queue.empty():
                try:
                    text = capture.queue.get(timeout=0.1)
                    yield f"data: {text}\n\n"
                except queue.Empty:
                    continue
            
            # Make sure thread is finished
            cli_thread.join()
            
            yield f"data: [DONE]\n\n"
            
        finally:
            # Restore original stdout
            sys.stdout = original_stdout
    
    return Response(generate(), mimetype='text/event-stream')

@app.route('/api/download/<research_id>/<format>', methods=['GET'])
def download_report(research_id, format):
    print(f"Download request: research_id={research_id}, format={format}")
    
    report = reports_db.get_report(research_id)
    if not report:
        print(f"Research ID {research_id} not found in database")
        return jsonify({'error': 'Research not found'}), 404
    
    print(f"Report found in database: {report.keys()}")
    
    if not report.get('content'):
        print(f"No content in report: {report}")
        return jsonify({'error': 'No content available for download'}), 400
    
    question = report['question']
    content = report['content']
    
    # Create a safe filename
    safe_title = "".join(c for c in question if c.isalnum() or c in (' ', '-', '_')).rstrip()[:50]
    
    if format == 'markdown':
        return Response(
            content,
            mimetype='text/markdown',
            headers={'Content-Disposition': f'attachment; filename="{safe_title}.md"'}
        )
    
    elif format == 'docx':
        docx_file = convert_to_docx(content, question)
        return send_file(
            docx_file,
            as_attachment=True,
            download_name=f"{safe_title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    elif format == 'pdf':
        pdf_file = convert_to_pdf(content, question)
        return send_file(
            pdf_file,
            as_attachment=True,
            download_name=f"{safe_title}.pdf",
            mimetype='application/pdf'
        )
    
    else:
        return jsonify({'error': 'Unsupported format'}), 400

@app.route('/api/reports/<research_id>', methods=['GET'])
def get_report(research_id):
    report = reports_db.get_report(research_id)
    if not report:
        return jsonify({'error': 'Research not found'}), 404
    
    return jsonify(report)

@app.route('/api/reports', methods=['GET'])
def list_reports():
    limit = request.args.get('limit', 50, type=int)
    offset = request.args.get('offset', 0, type=int)
    
    reports = reports_db.list_reports(limit=limit, offset=offset)
    return jsonify({
        'reports': reports,
        'limit': limit,
        'offset': offset
    })

@app.route('/api/reports/<research_id>', methods=['DELETE'])
def delete_report(research_id):
    success = reports_db.delete_report(research_id)
    if not success:
        return jsonify({'error': 'Research not found'}), 404
    
    return jsonify({'message': 'Report deleted successfully'})

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'deep_crawler_api'})

if __name__ == '__main__':
    host = CONFIG["server"]["api_host"]
    port = CONFIG["server"]["api_port"]
    print(f"Starting Deep Crawler API server on http://{host}:{port}")
    # Run without threading to avoid SQLite issues
    app.run(host=host, port=port, debug=False, threaded=False)
